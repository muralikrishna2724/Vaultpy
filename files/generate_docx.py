import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def main():
    doc = Document()

    # Title
    title = doc.add_heading('VaultPy Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1 Introduction
    doc.add_heading('1 Introduction', level=1)
    doc.add_paragraph(
        "Password management is an ongoing challenge, and remembering multiple strong passwords is often difficult. "
        "We wanted to build an application that runs locally, keeps data secure, and is easy to use without relying "
        "on a cloud service. Therefore, we created VaultPy, a Python-based password manager that stores everything securely "
        "on the local machine using reliable encryption instead of plain text."
    )

    # 1.1 Problem Statement
    doc.add_heading('1.1 Problem Statement', level=2)
    doc.add_paragraph(
        "Many people tend to reuse the same weak passwords across different accounts because remembering unique, strong "
        "passwords is too hard. While cloud-based password managers exist, some users prefer not to store their private "
        "data on remote servers. There is a clear need for a fast, straightforward, and local alternative that guarantees "
        "privacy while offering strong encryption to protect stored credentials."
    )

    # 1.2 Project Objectives
    doc.add_heading('1.2 Project Objectives', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("To build a simple, easy-to-use desktop application for managing passwords.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("To ensure maximum security by encrypting the data locally using AES-256-GCM.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("To protect the master password against cracking attempts using Argon2id hashing.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("To present a clean, modern user interface that behaves like a web app but runs natively on the desktop.")

    # 1.3 Software & Hardware specifications
    doc.add_heading('1.3 Software & Hardware specifications', level=2)

    # 1.3.1 Software requirements
    doc.add_heading('1.3.1 Software requirements', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Operating System: ").bold = True
    p.add_run("Windows 10 or higher")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Language: ").bold = True
    p.add_run("Python 3.x")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Frameworks/Libraries: ").bold = True
    p.add_run("Flask (for backend), PyWebView (for desktop UI), Cryptography (AES-256), Argon2-cffi (hashing)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Frontend: ").bold = True
    p.add_run("HTML, CSS, JavaScript (Vanilla)")

    # 1.3.2 Hardware requirements
    doc.add_heading('1.3.2 Hardware requirements', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Processor: ").bold = True
    p.add_run("Dual-core CPU or better (Intel i3/Ryzen 3 equivalent)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("RAM: ").bold = True
    p.add_run("Minimum 2GB (4GB recommended)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Storage: ").bold = True
    p.add_run("100 MB of free storage space for the application and the vault")

    # 2 Design Methodology
    doc.add_heading('2 Design Methodology', level=1)

    # 2.1 System Architecture
    doc.add_heading('2.1 System Architecture', level=2)
    doc.add_paragraph(
        "VaultPy functions as a local web application wrapper. We use a Flask backend that handles all data operations "
        "and cryptography. PyWebView serves as the frontend client, displaying the user interface naturally on the desktop. "
        "The backend communicates with a local encrypted file (vault.enc) to read and update user entries. Importantly, "
        "no data ever leaves the local machine."
    )

    # 2.2 Data flow Diagram or Flowchart
    doc.add_heading('2.2 Data flow Diagram or Flowchart', level=2)
    doc.add_paragraph(
        "For the data flow:\n"
        "1. The user launches the application and enters the Master Password.\n"
        "2. The system hashes the password using Argon2id to derive a secure 256-bit encryption key.\n"
        "3. This key is used to decrypt the vault.enc file using AES-256-GCM.\n"
        "4. If successful, the Flask server securely passes the decrypted entries to the PyWebView frontend.\n"
        "5. When a user adds a new entry, the frontend sends it to Flask, which updates the list in memory, "
        "re-encrypts the data, and saves the secure file on disk."
    )

    # 2.3 Technology Description
    doc.add_heading('2.3 Technology Description', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Python: ").bold = True
    p.add_run("The core programming language used to build the application logic.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Flask: ").bold = True
    p.add_run("Used to create the local web server that bridges the UI and the logic.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("PyWebView: ").bold = True
    p.add_run("Displays the Flask application in a standalone desktop window without a standard browser toolbar, "
              "giving it the feel of a native application.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Cryptography (AES-256-GCM): ").bold = True
    p.add_run("An authenticated encryption standard used to secure the vault configuration file.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Argon2id: ").bold = True
    p.add_run("A robust key derivation function to process the master password securely, protecting it "
              "against brute-force attacks.")

    # 3 Implementation & Testing
    doc.add_heading('3 Implementation & Testing', level=1)

    # 3.1 Code snippets
    doc.add_heading('3.1 Code snippets', level=2)
    doc.add_paragraph("Deriving the key using Argon2id:")
    snippet1 = doc.add_paragraph(
        'key = hash_secret_raw(\n'
        '    secret=master_password.encode("utf-8"),\n'
        '    salt=salt,\n'
        '    time_cost=3,\n'
        '    memory_cost=65536,\n'
        '    parallelism=4,\n'
        '    hash_len=32,\n'
        '    type=Type.ID\n'
        ')'
    )
    snippet1.style = 'Normal'
    run = snippet1.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph("Encrypting data using AES-GCM:")
    snippet2 = doc.add_paragraph(
        'aesgcm = AESGCM(key)\n'
        'nonce = os.urandom(12)\n'
        'ciphertext = aesgcm.encrypt(nonce, plaintext.encode("utf-8"), None)\n'
    )
    snippet2.style = 'Normal'
    run = snippet2.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # 3.2 Test cases
    doc.add_heading('3.2 Test cases', level=2)
    doc.add_paragraph("Test Case 1: Entering the correct master password.\nExpected: Vault unlocks and shows dashboard.\nResult: Pass.", style='List Bullet')
    doc.add_paragraph("Test Case 2: Entering an incorrect master password.\nExpected: Application rejects access and shows an error message.\nResult: Pass.", style='List Bullet')
    doc.add_paragraph("Test Case 3: Adding a new password entry.\nExpected: Entry is displayed in the list and saved encrypted on the local disk.\nResult: Pass.", style='List Bullet')
    doc.add_paragraph("Test Case 4: Copying a password.\nExpected: Password copies to the clipboard and clears out automatically after 30 seconds for security purposes.\nResult: Pass.", style='List Bullet')

    # 4 Conclusion
    doc.add_heading('4 Conclusion', level=1)
    doc.add_paragraph(
        "In conclusion, VaultPy successfully meets our goal of building a lightweight, secure, and fully local password manager. "
        "Working on this project allowed us to practically learn about cybersecurity and cryptography, specifically implementing "
        "AES-256-GCM and Argon2id. It is a functional application that addresses the problem of managing passwords safely "
        "without needing to rely on third-party cloud tools."
    )

    # Bibliography
    doc.add_heading('Bibliography', level=1)
    doc.add_paragraph("Python Official Documentation: https://docs.python.org/3/", style='List Bullet')
    doc.add_paragraph("Cryptography PyPI Documentation: https://cryptography.io/en/latest/", style='List Bullet')
    doc.add_paragraph("Flask Web Development guide", style='List Bullet')
    doc.add_paragraph("Argon2 documentation and security guidelines", style='List Bullet')

    doc.add_page_break()

    # Appendix: (Source code)
    doc.add_heading('Appendix: (Source code)', level=1)
    doc.add_paragraph("These are the main files required to set up the VaultPy application logic and backend.")

    # List of files to add to appendix
    files_to_add = ['app.py', 'crypto.py', 'vault_manager.py', 'requirements.txt']
    
    for filename in files_to_add:
        filepath = os.path.join(BASE_DIR, filename)
        if os.path.exists(filepath):
            doc.add_heading(filename, level=2)
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            p = doc.add_paragraph(content)
            p.style = 'Normal'
            # Change font to Consolas for code
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Consolas'
                run.font.size = Pt(8)

    save_path = os.path.join(BASE_DIR, 'VaultPy_Documentation.docx')
    doc.save(save_path)
    print(f"Documentation saved to {save_path}")

    pdf_path = os.path.join(BASE_DIR, 'VaultPy_Documentation.pdf')
    convert(save_path, pdf_path)
    print(f"PDF saved to {pdf_path}")

if __name__ == "__main__":
    main()
